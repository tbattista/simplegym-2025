from docx import Document
from pathlib import Path
import tempfile
import os
from datetime import datetime
from typing import Dict, Any
from ..models import WorkoutData
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("Warning: docx2pdf not available. PDF generation will be disabled.")

class DocumentService:
    """Service for processing Word documents and replacing template variables"""
    
    def __init__(self):
        self.temp_dir = Path("backend/uploads")
        self.temp_dir.mkdir(exist_ok=True)
    
    def generate_document(self, workout_data: WorkoutData, template_path: Path) -> Path:
        """
        Generate a filled Word document from template and workout data
        
        Args:
            workout_data: The workout information to fill into the template
            template_path: Path to the template Word document
            
        Returns:
            Path to the generated document file
        """
        try:
            # Load the template document
            doc = Document(template_path)
            
            # Create replacement dictionary
            replacements = self._create_replacements(workout_data)
            
            # Replace variables in the document
            self._replace_variables_in_document(doc, replacements)
            
            # Generate output filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"gym_log_{workout_data.workout_name.replace(' ', '_')}_{timestamp}.docx"
            output_path = self.temp_dir / output_filename
            
            # Save the modified document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error generating document: {str(e)}")
    
    def _create_replacements(self, workout_data: WorkoutData) -> Dict[str, str]:
        """
        Create a dictionary of all template variables and their replacements
        
        Args:
            workout_data: The workout data containing all the information
            
        Returns:
            Dictionary mapping template variables to replacement values
        """
        replacements = {
            # Basic workout information with proper template format
            '{{ workout_name }}': workout_data.workout_name,
            "today's date:": f"today's date: {workout_data.workout_date}",
        }
        
        # Add all exercise replacements with proper template format (double braces)
        for key, value in workout_data.exercises.items():
            replacements[f'{{{{ {key} }}}}'] = value
        
        # Add sets, reps, and rest replacements
        for key, value in workout_data.sets.items():
            replacements[f'{{{{ {key} }}}}'] = value
            
        for key, value in workout_data.reps.items():
            replacements[f'{{{{ {key} }}}}'] = value
            
        for key, value in workout_data.rest.items():
            replacements[f'{{{{ {key} }}}}'] = value
        
        # Add bonus exercise replacements with proper template format
        for key, value in workout_data.bonus_exercises.items():
            replacements[f'{{{{ {key} }}}}'] = value
            
        # Add bonus sets, reps, and rest replacements
        for key, value in workout_data.bonus_sets.items():
            replacements[f'{{{{ {key} }}}}'] = value
            
        for key, value in workout_data.bonus_reps.items():
            replacements[f'{{{{ {key} }}}}'] = value
            
        for key, value in workout_data.bonus_rest.items():
            replacements[f'{{{{ {key} }}}}'] = value
        
        return replacements
    
    def _replace_variables_in_document(self, doc: Document, replacements: Dict[str, str]) -> None:
        """
        Replace template variables in the Word document
        
        Args:
            doc: The Word document object
            replacements: Dictionary of variables to replace
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_text(paragraph, replacements)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_text(paragraph, replacements)
    
    def _replace_in_text(self, paragraph, replacements: Dict[str, str]) -> None:
        """
        Replace variables in a paragraph while preserving formatting
        
        Args:
            paragraph: The paragraph object to process
            replacements: Dictionary of variables to replace
        """
        # Get the full text of the paragraph
        full_text = paragraph.text
        
        # Check if any replacements are needed or if there are template variables
        import re
        has_template_vars = bool(re.search(r'\{\{.*?\}\}', full_text))
        has_replacements = any(find_text in full_text for find_text in replacements.keys())
        
        if not has_template_vars and not has_replacements:
            return
        
        # Perform replacements
        new_text = full_text
        for find_text, replace_text in replacements.items():
            if find_text in new_text:
                new_text = new_text.replace(find_text, replace_text)
        
        # Remove any remaining template variables (including partial ones)
        # This handles {{ variable }}, { variable }, and leftover braces
        new_text = re.sub(r'\{\{\s*[^}]*\s*\}\}', '', new_text)  # Remove {{ ... }}
        new_text = re.sub(r'\{\s*[^}]*\s*\}', '', new_text)      # Remove { ... }
        
        # Clean up any extra spaces that might be left
        new_text = re.sub(r'\s+', ' ', new_text).strip()
        
        # If text changed, update the paragraph
        if new_text != full_text:
            # Clear existing runs and add new text
            paragraph.clear()
            paragraph.add_run(new_text)
    
    def get_template_variables(self, template_path: Path) -> Dict[str, Any]:
        """
        Extract template variables from a Word document
        
        Args:
            template_path: Path to the template document
            
        Returns:
            Dictionary containing information about template variables
        """
        try:
            doc = Document(template_path)
            
            # Extract all text content
            all_text = ""
            
            # Get text from paragraphs
            for paragraph in doc.paragraphs:
                all_text += paragraph.text + "\n"
            
            # Get text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text + " "
                    all_text += "\n"
            
            # Find template variables (assuming {{ variable }} format)
            import re
            variables = re.findall(r'\{\{\s*([^}]+)\s*\}\}', all_text)
            
            # Also look for exercise- pattern variables
            exercise_vars = re.findall(r'exercise-\w+', all_text)
            
            return {
                "template_variables": list(set(variables)),
                "exercise_variables": list(set(exercise_vars)),
                "content_preview": all_text[:500] + "..." if len(all_text) > 500 else all_text
            }
            
        except Exception as e:
            raise Exception(f"Error analyzing template: {str(e)}")
    
    def generate_preview_pdf(self, workout_data: WorkoutData, template_path: Path) -> Path:
        """
        Generate a PDF preview of the filled document
        
        Args:
            workout_data: The workout information to fill into the template
            template_path: Path to the template Word document
            
        Returns:
            Path to the generated PDF file
        """
        if not DOCX2PDF_AVAILABLE:
            raise Exception("PDF generation is not available on this server. Please download the Word document instead.")
        
        try:
            # First generate the Word document
            word_path = self.generate_document(workout_data, template_path)
            
            # Convert to PDF
            pdf_path = self._convert_to_pdf(word_path)
            
            return pdf_path
            
        except Exception as e:
            raise Exception(f"Error generating PDF preview: {str(e)}")
    
    def _convert_to_pdf(self, word_path: Path) -> Path:
        """
        Convert a Word document to PDF
        
        Args:
            word_path: Path to the Word document
            
        Returns:
            Path to the generated PDF file
        """
        if not DOCX2PDF_AVAILABLE:
            raise Exception("PDF conversion is not available on this server.")
        
        try:
            # Generate PDF filename
            pdf_path = word_path.with_suffix('.pdf')
            
            # Convert Word to PDF
            convert(str(word_path), str(pdf_path))
            
            return pdf_path
            
        except Exception as e:
            raise Exception(f"Error converting to PDF: {str(e)}")
    
    def cleanup_old_files(self, max_age_hours: int = 24) -> int:
        """
        Clean up old generated files
        
        Args:
            max_age_hours: Maximum age of files to keep in hours
            
        Returns:
            Number of files deleted
        """
        try:
            deleted_count = 0
            current_time = datetime.now()
            
            # Clean up both Word and PDF files
            for pattern in ["gym_log_*.docx", "gym_log_*.pdf"]:
                for file_path in self.temp_dir.glob(pattern):
                    # Get file modification time
                    file_time = datetime.fromtimestamp(file_path.stat().st_mtime)
                    age_hours = (current_time - file_time).total_seconds() / 3600
                    
                    if age_hours > max_age_hours:
                        file_path.unlink()
                        deleted_count += 1
            
            return deleted_count
            
        except Exception as e:
            # Log error but don't fail the application
            print(f"Warning: Error cleaning up old files: {str(e)}")
            return 0
